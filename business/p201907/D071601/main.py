import docx
import pandas as pd
from collections import Counter, OrderedDict
import wordcloud
import matplotlib.pyplot as plt
import os
from jieba import analyse
import re
import random

"""
词频分析, 高频词对比分析
共现矩阵
聚类分析
主题模型
"""

# 画图 显示中文
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

# 创建结果文件夹
if not os.path.exists("./results"):
    os.makedirs("./results")

# 引入TextRank关键词抽取接口
textrank = analyse.textrank


def get_kw(text):
    out = []
    keywords = textrank(text, topK=random.choice(range(5, 20)))
    for keyword in keywords:
        out.append(keyword)
    return out


def readLines(filename):
    # 读文件的方法
    out = []
    with open(filename, 'r', encoding='utf-8') as fr:
        lines = fr.readlines()
        for line in lines:
            out.append(line.replace("\n", ""))
    return out


stopwords = readLines('stopwords.txt')


def plot_word_cloud(data_name, max_words):
    text = open('./results/{}.txt'.format(data_name), 'r', encoding='utf-8').read()

    wc = wordcloud.WordCloud(width=200, height=180, background_color="white",
                             max_font_size=60,
                             random_state=1,
                             max_words=max_words,
                             font_path='C:\\Windows\\msyh.ttf', repeat=True)
    wc.generate(text)
    plt.axis("off")
    plt.figure()
    wc.to_file('./results/{}_word_cloud.png'.format(data_name))


def build_data():
    """
    利用正则表达式处理每年的数据，把每个报道的年份，月份提取出来，最后保存到dataAll.xlsx文件中
    数据格式：
            年份，月份，日期，内容
            --，--，--，--，--

    :return:
    """
    save = []
    # 处理13,14,15,16年的数据
    for y in ['2013', '2014', '2015', '2016']:
        for file in os.listdir("./{}年扶贫报道".format(y)):
            file_name = os.path.join("./{}年扶贫报道".format(y), file)
            with open(file_name, 'r', encoding='gbk') as fr:
                lines = fr.readlines()
                one_port = []
                for line in lines:
                    one_port.append(line.replace("\n", ""))
                publish_date = re.search("<日期>=\d+\.\d+\.\d+", "".join(one_port))
                if publish_date:
                    dates = publish_date.group().replace("<日期>=", '').replace("{}.".format(y), '')
                    rows = OrderedDict()
                    rows['年份'] = '{}'.format(y)
                    rows['month'] = dates.split(".")[0]
                    rows['day'] = dates.split(".")[1]
                    rows['content'] = "".join(one_port)
                    save.append(rows)
    # 处理17年的数据
    data2017 = []
    with open("2017扶贫报道.txt", 'r', encoding='gbk') as fr:
        lines = fr.readlines()
        for line in lines:
            data2017.append(line.replace('\n', ''))
    data2017_sp = "".join(data2017).split("版")
    length = len(data2017_sp)

    for p in range(length):
        if p + 1 < length:
            one = data2017_sp[p + 1]
            publish_date2017 = re.search("【人民日报\d+年\d+月\d+日 第\d+", data2017_sp[p])
            if publish_date2017:
                dates = publish_date2017.group().replace("人民日报", '')
                rows2017 = OrderedDict()
                rows2017['年份'] = re.search("\d+年", dates).group().replace('年', '')
                rows2017['month'] = re.search("\d+月", dates).group().replace('月', '')
                rows2017['day'] = re.search("\d+日", dates).group().replace('日', '')
                rows2017['content'] = one
                save.append(rows2017)
    # 处理18,19年的数据，
    data1819 = []
    for file in os.listdir("./人民日报扶贫报道2018-2019"):
        file_name = os.path.join('./人民日报扶贫报道2018-2019', file)
        doc_file = docx.Document(file_name)
        for paragraph in doc_file.paragraphs:
            data1819.append(paragraph.text.replace(" ", "").replace("\n", ""))

    for one in re.split("版", "".join(data1819)):
        publish_date1819 = re.search("\d+年\d+月\d+日\d+", one)
        if publish_date1819:
            dates = publish_date1819.group()
            rows1819 = OrderedDict()
            rows1819['年份'] = re.search("\d+年", dates).group().replace('年', '')
            rows1819['month'] = re.search("\d+月", dates).group().replace('月', '')
            rows1819['day'] = re.search("\d+日", dates).group().replace('日', '')
            rows1819['content'] = one
            save.append(rows1819)

    data = pd.DataFrame(save)
    data.to_excel("dataAll.xlsx", index=None)
    print(data.shape)


def split_time(y, m):
    """
        2013年2月-2014年2月==1
    2014年2月-2015年11月==2
    2015年12月-2016年11月==3
    2016年12月-2017年12月==4
    2018年1月-2019年5月==5
    :return:
    """
    if y == 2013:
        return '1'
    if y == 2014 and m < 3:
        return '1'
    if y == 2014 and m >= 3:
        return '2'
    if y == 2015 and m <= 11:
        return '2'
    if y == 2015 and m > 11:
        return '3'
    if y == 2016 and m <= 11:
        return '3'
    if y == 2016 and m > 11:
        return '4'
    if y == 2017 and m <= 12:
        return '4'
    if y == 2018 or y == 2019:
        return '5'


def cutData():
    # 分词方法
    part_1 = open('./results/part_1.txt', 'w', encoding='utf-8')
    part_2 = open('./results/part_2.txt', 'w', encoding='utf-8')
    part_3 = open('./results/part_3.txt', 'w', encoding='utf-8')
    part_4 = open('./results/part_4.txt', 'w', encoding='utf-8')
    part_5 = open('./results/part_5.txt', 'w', encoding='utf-8')

    data = pd.read_excel("dataAll.xlsx")
    data['label'] = data.apply(lambda row: split_time(row['年份'], row['month']), axis=1)

    for x, y in data.groupby(by='label'):
        print("xxxx:{}".format(x))
        paragraph = y['content'].values.tolist()
        words = []
        for p in paragraph:
            p_cut = get_kw(p)
            for i in p_cut:
                words.append(i)
            if x == '1':
                part_1.writelines(" ".join(p_cut) + "\n")
            if x == '2':
                part_2.writelines(" ".join(p_cut) + "\n")
            if x == '3':
                part_3.writelines(" ".join(p_cut) + "\n")
            if x == '4':
                part_4.writelines(" ".join(p_cut) + "\n")
            if x == '5':
                part_5.writelines(" ".join(p_cut) + "\n")

        words_freq = []
        # 输出频率top 50 的高频词到excel文件
        for wd, freq in Counter(words).most_common(50):
            rows = OrderedDict()
            rows['word'] = wd
            rows['freq'] = freq
            words_freq.append(rows)

        df = pd.DataFrame(words_freq)
        df.to_excel("./results/words_freq_part_{}.xlsx".format(x), index=None)


if __name__ == '__main__':
    method = 'cluster'

    if method == 'build_data':
        # 把每年的数据组合起来，保存到一个文件里面
        build_data()

    if method == 'cutData':
        # 对数据进行分词，提取关键词
        cutData()

    if method == 'plot_word_cloud':
        # 做词云图
        max_words = 20
        plot_word_cloud(data_name='part_1', max_words=max_words)
        plot_word_cloud(data_name='part_2', max_words=max_words)
        plot_word_cloud(data_name='part_3', max_words=max_words)
        plot_word_cloud(data_name='part_4', max_words=max_words)
        plot_word_cloud(data_name='part_5', max_words=max_words)

    if method == 'getoccurrences':
        # 共现矩阵分析
        token_sent_list = []
        all_path = ['./results/part_{}.txt'.format(i) for i in range(1, 6)]

        for y in all_path:
            data2018 = []
            with open(y, 'r', encoding='utf-8') as fr:
                lines = fr.readlines()
                for line in lines:
                    for words in line.split():
                        data2018.append(words)

            token_sent_list.append(data2018)

        vocab = list(set([i for j in token_sent_list for i in j]))
        print("vocab:{}".format(len(vocab)))
        co_occ = {ii: Counter({jj: 0 for jj in vocab if jj != ii}) for ii in vocab}
        k = 2

        for sen in token_sent_list:
            for ii in range(len(sen)):
                if ii < k:
                    c = Counter(sen[0:ii + k + 1])
                    del c[sen[ii]]
                    co_occ[sen[ii]] = co_occ[sen[ii]] + c
                elif ii > len(sen) - (k + 1):
                    c = Counter(sen[ii - k::])
                    del c[sen[ii]]
                    co_occ[sen[ii]] = co_occ[sen[ii]] + c
                else:
                    c = Counter(sen[ii - k:ii + k + 1])
                    del c[sen[ii]]
                    co_occ[sen[ii]] = co_occ[sen[ii]] + c

        co_occ = {ii: dict(co_occ[ii]) for ii in vocab}
        df = pd.DataFrame(co_occ)
        df.to_excel("./results/共现矩阵.xlsx")
        print(df.shape)

    if method == 'LdaModel':
        # LDA 主题模型，输出文件ldaOutput.txt 表示输出的文档主题。
        from gensim.corpora import Dictionary
        from gensim.models import LdaModel
        from gensim import models

        # LDA 主题模型
        # 构建训练语料
        all_path = ['./results/part_{}.txt'.format(i) for i in range(1, 6)]

        for y in all_path:
            train_set = []
            with open(y, 'r', encoding='utf-8') as fr:
                lines = fr.readlines()
                for line in lines:
                    train_set.append([i for i in line.split()])
            print(len(train_set), train_set[0])

            dictionary = Dictionary(train_set)
            corpus = [dictionary.doc2bow(text) for text in train_set]  # 构建稀疏向量
            tfidf = models.TfidfModel(corpus)  # 统计tfidf
            corpus_tfidf = tfidf[corpus]  # 得到每个文本的tfidf向量，稀疏矩阵
            lda = LdaModel(corpus_tfidf, id2word=dictionary, num_topics=50, iterations=10)
            test = lda.print_topics(50)
            fw = open('./results/lda_{}'.format(y.replace("./results/", "")), 'w', encoding='utf-8')
            for i in test:
                print("{} {}\n".format(i[0], i[1]))
                fw.writelines("{} {}\n".format(i[0], i[1]))

    if method == 'cluster':
        # 聚类分析
        from gensim.models.doc2vec import Doc2Vec
        from sklearn.cluster import KMeans
        import gensim

        TaggededDocument = gensim.models.doc2vec.TaggedDocument

        # 利用gensim将doc转换为向量
        all_path = ['./results/part_{}.txt'.format(i) for i in range(1, 6)]
        train_set = []

        for y in all_path:
            with open(y, 'r', encoding='utf-8') as fr:
                lines = fr.readlines()
                for line in lines:
                    train_set.append([i for i in line.split()])

        x_train = []
        for i, word_list in enumerate(train_set):
            # 遍历每一条评论
            document = TaggededDocument(word_list, tags=[i])
            x_train.append(document)

        print("document length : {}".format(len(x_train)))
        # 训练词向量，大小是100维， 这里维度可以调整
        model_dm = Doc2Vec(x_train, min_count=1, window=3, vector_size=100, sample=1e-3, negative=5, workers=4)
        model_dm.train(x_train, total_examples=model_dm.corpus_count, epochs=100)
        model_dm.save('model_dm.model')  # 保存模型

        infered_vectors_list = []
        i = 0
        for text, label in x_train:
            vector = model_dm.infer_vector(text)  # 计算指定句子的向量
            infered_vectors_list.append(vector)
            i += 1

        print("infered_vectors_list length :{}".format(len(infered_vectors_list)))
        n_cluster = 10

        print("============= Kmeans 聚类 =============")
        kmean_model = KMeans(n_clusters=n_cluster, init='k-means++', n_init=10,
                             max_iter=1000, tol=1e-4, precompute_distances='auto',
                             verbose=0, random_state=None, copy_x=True,
                             n_jobs=None, algorithm='auto')
        kmean_model.fit(infered_vectors_list)
        cluster_label = kmean_model.labels_

        df = pd.DataFrame()
        df['data'] = x_train
        df['label'] = cluster_label
        df.to_excel("聚类结果.xlsx", index=None)
        print(df.shape)
