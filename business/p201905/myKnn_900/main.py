from sklearn.neighbors import KNeighborsClassifier
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, classification_report
from sklearn.externals import joblib
import numpy as np
import docx
import os
import pandas as pd

stopwords = []
doc_file = docx.Document('暂停词或介词.docx')
for paragraph in doc_file.paragraphs:
    stopwords.append(paragraph.text.strip())

fumian = []
doc_file = docx.Document('负面词.docx')
for paragraph in doc_file.paragraphs:
    fumian.append(paragraph.text.strip())


def get_label(inputs):
    for i in fumian:
        if i in inputs.split():
            # 负面
            return 0
    # 正面
    return 1


def train():
    text_cut_train = []
    text_cut_label = []
    save_data = []
    for file in os.listdir("./data"):
        file_name = os.path.join('./data', file)
        doc_file = docx.Document(file_name)
        k = 0
        for paragraph in doc_file.paragraphs:
            labels = get_label(paragraph.text)
            rows = {'data': paragraph.text, 'labels': labels}
            save_data.append(rows)
            k += 1
            text_cut = [i for i in paragraph.text.split() if i not in stopwords]
            text_cut_train.append(" ".join(text_cut))
            text_cut_label.append(labels)

    df_data = pd.DataFrame(save_data)
    df_data.to_excel("save_data.xlsx", index=None)

    tfidf2 = TfidfVectorizer()
    data = tfidf2.fit_transform(text_cut_train)
    df = pd.DataFrame(data.toarray())
    df.to_excel("train.xlsx", index=None)

    x_train, x_test, y_train, y_test = train_test_split(data, np.array(text_cut_label), test_size=0.3)
    print(x_train.shape, x_test.shape, y_train.shape, y_test.shape)

    model = KNeighborsClassifier(n_neighbors=5, weights='uniform', algorithm='auto', leaf_size=30, p=2,
                                 metric='minkowski')
    model.fit(x_train, y_train)
    joblib.dump(model, 'model.m')
    y_test_ = model.predict(x_test)
    acc = accuracy_score(y_test, y_test_)
    print('准确率：')
    print("{:0.2f}".format(acc))
    print(classification_report(y_test, y_test_))


def train_2():
    text_cut_train = []
    text_cut_label = []
    save_data = []
    title_data = []

    for file in os.listdir("./xlsx"):
        data = pd.read_excel(os.path.join('./xlsx', file))
        for x, y in data.iterrows():
            paragraph = y['data']
            labels = y['labels']

            title = y['标题']
            title_data.append(title)
            rows = {'data': paragraph, 'labels': labels}
            save_data.append(rows)
            text_cut = [i for i in paragraph.split() if i not in stopwords]
            text_cut_train.append(" ".join(text_cut))
            text_cut_label.append(labels)

    tfidf2 = TfidfVectorizer()
    data = tfidf2.fit_transform(text_cut_train)

    x_train, x_test, y_train, y_test = train_test_split(data, np.array(text_cut_label), test_size=0.3)
    print(x_train.shape, x_test.shape, y_train.shape, y_test.shape)

    model = KNeighborsClassifier(n_neighbors=5, weights='uniform', algorithm='auto', leaf_size=30, p=2,
                                 metric='minkowski')
    model.fit(x_train, y_train)
    joblib.dump(model, 'model.m')
    y_test_ = model.predict(x_test)
    acc = accuracy_score(y_test, y_test_)
    print('准确率：')
    print("{:0.2f}".format(acc))
    print(classification_report(y_test, y_test_))
    # 预测
    y_test_ = model.predict(data)
    df = pd.DataFrame()
    df['text_cut_label'] = text_cut_label
    df['data'] = text_cut_train
    df['predict'] = y_test_
    df['title'] = title_data
    df.to_excel("KNN预测结果.xlsx", index=None)
    print(df.shape)
    print("预测完成")


train_2()
